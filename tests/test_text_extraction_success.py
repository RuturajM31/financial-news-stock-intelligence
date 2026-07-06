"""Successful text-extraction tests."""

import csv

from docx import Document

from financial_news_intelligence.data import text_extraction


# ============================================================
# 1. PASTED TEXT
# ============================================================

def test_extract_from_text_cleans_whitespace() -> None:
    """Pasted text should return one clean sentence."""

    # Prepare messy article text with spaces and a line break.
    article_text = "  NVIDIA   reported\nrecord revenue.  "

    # Run the pasted-text extraction function.
    result = text_extraction.extract_from_text(article_text)

    # Confirm repeated whitespace was replaced with single spaces.
    assert result == "NVIDIA reported record revenue."


# ============================================================
# 2. TXT FILE
# ============================================================

def test_extract_from_txt(tmp_path) -> None:
    """TXT extraction should read and clean the file."""

    # Create a temporary TXT file used only by this test.
    file_path = tmp_path / "article.txt"

    # Add messy spacing so the cleaner has something to fix.
    file_path.write_text(
        "Apple   reported higher profit.",
        encoding="utf-8",
    )

    # Run the real TXT extraction function.
    result = text_extraction.extract_from_txt(file_path)

    # Confirm the returned article contains clean spacing.
    assert result == "Apple reported higher profit."


# ============================================================
# 3. PDF FILE
# ============================================================

def test_extract_from_pdf(monkeypatch, tmp_path) -> None:
    """PDF extraction should combine text from all pages."""

    # Create a temporary PDF path so file validation succeeds.
    file_path = tmp_path / "article.pdf"
    file_path.write_bytes(b"temporary test file")

    class FakePage:
        """Small fake PDF page used by the test."""

        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            # Return controlled page text instead of reading a real PDF.
            return self.text

    class FakePdfReader:
        """Small fake PDF reader with two test pages."""

        def __init__(self, _file_path) -> None:
            self.pages = [
                FakePage("Tesla reported"),
                FakePage("higher deliveries."),
            ]

    # Replace the real PDF reader with the controlled test reader.
    monkeypatch.setattr(
        text_extraction,
        "PdfReader",
        FakePdfReader,
    )

    # Run the real PDF extraction function.
    result = text_extraction.extract_from_pdf(file_path)

    # Confirm text from both pages was joined and cleaned.
    assert result == "Tesla reported higher deliveries."


# ============================================================
# 4. DOCX FILE
# ============================================================

def test_extract_from_docx(tmp_path) -> None:
    """DOCX extraction should combine readable paragraphs."""

    # Create a temporary DOCX file used only by this test.
    file_path = tmp_path / "article.docx"

    # Build a small document with two article paragraphs.
    document = Document()
    document.add_paragraph("Microsoft reported")
    document.add_paragraph("strong cloud growth.")
    document.save(file_path)

    # Run the real DOCX extraction function.
    result = text_extraction.extract_from_docx(file_path)

    # Confirm both paragraphs were joined into one clean article.
    assert result == "Microsoft reported strong cloud growth."


# ============================================================
# 5. CSV FILE
# ============================================================

def test_extract_from_csv(tmp_path) -> None:
    """CSV extraction should return one article per valid row."""

    # Create a temporary CSV file used only by this test.
    file_path = tmp_path / "articles.csv"

    # Write two financial-news rows into the article_text column.
    with file_path.open("w", encoding="utf-8", newline="") as file:
        writer = csv.DictWriter(
            file,
            fieldnames=["headline", "article_text"],
        )
        writer.writeheader()
        writer.writerow(
            {
                "headline": "News 1",
                "article_text": "Amazon   raised its forecast.",
            }
        )
        writer.writerow(
            {
                "headline": "News 2",
                "article_text": "Meta reported lower costs.",
            }
        )

    # Run the CSV extraction function using the selected text column.
    result = text_extraction.extract_from_csv(
        file_path,
        text_column="article_text",
    )

    # Confirm each CSV row became one clean article.
    assert result == [
        "Amazon raised its forecast.",
        "Meta reported lower costs.",
    ]


# ============================================================
# 6. URL
# ============================================================

def test_extract_from_url(monkeypatch) -> None:
    """URL extraction should collect visible paragraph text."""

    class FakeResponse:
        """Controlled webpage response used instead of the internet."""

        text = """
        <html>
            <body>
                <nav>Navigation</nav>
                <p>NVIDIA reported record revenue.</p>
                <p>The company raised its outlook.</p>
            </body>
        </html>
        """

        def raise_for_status(self) -> None:
            # Simulate a successful web request.
            return None

    def fake_get(*_args, **_kwargs):
        # Return controlled HTML instead of calling a real website.
        return FakeResponse()

    # Replace requests.get so this test stays fast and offline.
    monkeypatch.setattr(
        text_extraction.requests,
        "get",
        fake_get,
    )

    # Run the real URL extraction function.
    result = text_extraction.extract_from_url(
        "https://example.com/article"
    )

    # Confirm navigation was removed and paragraphs were joined.
    assert result == (
        "NVIDIA reported record revenue. "
        "The company raised its outlook."
    )
